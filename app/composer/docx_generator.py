from docx import Document


def create_docx(filename, title, content):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    doc.save(filename)


if __name__ == '__main__':
    create_docx('example.docx', 'Document Title', 'This is the content of the document.')